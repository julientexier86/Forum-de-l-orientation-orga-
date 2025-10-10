# 📄 2_Affectation_Eleves.py

import streamlit as st
import pandas as pd
from datetime import timedelta
from io import BytesIO
import os

# 🔐 Titre
st.title("📚 Forum Orientation - Traitement des affectations")
st.markdown("""
Cette page vous permet de :
- Charger les fichiers `voeux`, `tables`, `groupes`
- Générer les plannings des élèves (Excel, Word, PDF)
- Produire les plannings par intervenant

Toutes les affectations respectent les vœux des élèves et les créneaux disponibles.
""")

# ✨ Upload des fichiers
col1, col2, col3 = st.columns(3)

with col1:
    voeux_file = st.file_uploader("1. Vœux des élèves", type="xlsx")
with col2:
    tables_file = st.file_uploader("2. Tables / Métiers", type="xlsx")
with col3:
    groupes_file = st.file_uploader("3. Groupes horaires", type="xlsx")

if voeux_file and tables_file and groupes_file:

    # 📂 Lecture des fichiers
    df_voeux = pd.read_excel(voeux_file)
    df_tables = pd.read_excel(tables_file)
    df_groupes = pd.read_excel(groupes_file)

    # Nettoyage
    df_tables.columns = df_tables.columns.str.strip()
    df_groupes.columns = df_groupes.columns.str.strip()

    if st.button("🚀 Générer les plannings"):
        # ⏰ Créneaux horaires par classe
        groupes_horaires = {}
        for _, row in df_groupes.iterrows():
            groupe = row["Groupe"]
            debut = pd.to_datetime(row["Horaire début"], format="%H:%M")
            fin = pd.to_datetime(row["Horaire fin"], format="%H:%M")

            creneaux = []
            current = debut
            while current + timedelta(minutes=15) <= fin:
                creneaux.append(current.strftime("%H:%M"))
                current += timedelta(minutes=15)

            groupes_horaires[groupe] = creneaux

        # 🗓️ Agenda des tables
        agenda_tables = {}
        for _, row in df_tables.iterrows():
            capacite = row["Capacite par creneau"]
            debut = pd.to_datetime(row["Heure debut"], format="%H:%M")
            fin = pd.to_datetime(row["Heure fin"], format="%H:%M")
            current = debut
            while current + timedelta(minutes=15) <= fin:
                heure = current.strftime("%H:%M")
                agenda_tables[(row["Metier"], heure)] = capacite
                current += timedelta(minutes=15)

        # 📊 Affectations
        affectations = []
        for _, eleve in df_voeux.iterrows():
            nom = eleve["Nom"]
            prenom = eleve["Prénom"]
            classe = eleve["Classe"]
            creneaux = groupes_horaires.get(classe, [])
            voeux = [eleve[f"Vœu {i}"] for i in range(1, 7)]
            dejas = set()
            affecte_creneaux = set()

            for vœu in voeux:
                for heure in creneaux:
                    key = (vœu, heure)
                    if (key in agenda_tables
                        and agenda_tables[key] > 0
                        and vœu not in dejas
                        and heure not in affecte_creneaux):

                        affectations.append({
                            "Nom": nom,
                            "Prénom": prenom,
                            "Classe": classe,
                            "Heure": heure,
                            "Métier": vœu,
                            "Vœu n°": voeux.index(vœu) + 1
                        })
                        agenda_tables[key] -= 1
                        dejas.add(vœu)
                        affecte_creneaux.add(heure)
                        break
                if len(affecte_creneaux) >= 4:
                    break

        df_affect = pd.DataFrame(affectations)

        # Sauvegarde pour les autres pages
        os.makedirs("data", exist_ok=True)
        df_affect.to_csv("data/df_affectations.csv", index=False)
        st.session_state["df_affect"] = df_affect
        df_tables.to_csv("data/df_tables.csv", index=False)

        st.success(f"{len(df_affect)} affectations réalisées")

        # 👤 Affichage preview
        st.subheader("🎓 Aperçu des affectations")
        st.dataframe(df_affect.head(15))

        # 💾 Export Excel
        def to_excel(df):
            output = BytesIO()
            with pd.ExcelWriter(output, engine='openpyxl') as writer:
                df.to_excel(writer, index=False)
            output.seek(0)
            return output

        st.download_button("🔗 Télécharger les affectations (.xlsx)", data=to_excel(df_affect), file_name="affectations_forum.xlsx")

        # 📄 Génération d'un document Word avec planning par intervenant
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        from docx.oxml.ns import qn
        from docx.shared import Inches
        from docx.enum.style import WD_STYLE_TYPE

        metier_to_intervenant = dict(zip(df_tables["Metier"], df_tables["Nom Intervenant"]))
        df_affect["Intervenant"] = df_affect["Métier"].map(metier_to_intervenant)

        doc = Document()
        doc.add_heading('Planning des Intervenants – Forum Orientation', 0)

        df_sorted_interv = df_affect.sort_values(by=["Intervenant", "Heure"])

        for intervenant, df_int in df_sorted_interv.groupby("Intervenant"):
            if doc.paragraphs:
                doc.add_page_break()

            doc.add_heading(f'Planning de {intervenant}', level=1)
            df_int = df_int.sort_values(by="Heure")

            current_heure = None
            for _, row in df_int.iterrows():
                heure = row["Heure"]
                if heure != current_heure:
                    doc.add_paragraph("")  # Espace visuel
                    doc.add_paragraph(f"⏰ {heure}", style='Intense Quote')
                    current_heure = heure

                ligne = f"- {row['Nom']} {row['Prénom']} ({row['Classe']}) – {row['Métier']}"
                doc.add_paragraph(ligne, style='List Bullet')

        word_filename = "planning_intervenants.docx"
        doc.save(word_filename)

        with open(word_filename, "rb") as f:
            st.download_button("📥 Télécharger le planning par intervenant (.docx)", data=f, file_name=word_filename, mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    else:
        st.info("Appuie sur le bouton ci-dessus pour lancer l’affectation automatique.")

else:
    st.session_state.pop("df_affect", None)
    st.warning("Merci d'importer les 3 fichiers .xlsx pour démarrer.")
