import streamlit as st
import pandas as pd
import re, os, io, zipfile
from datetime import timedelta
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

st.set_page_config(page_title="Forum Orientation", page_icon="🎓", layout="wide")

# ==========================================
# FONCTIONS UTILITAIRES DE L'ALGORITHME
# ==========================================

def clean_cols(df):
    df = df.copy()
    df.columns = (
        df.columns.astype(str)
        .str.replace("\u00A0", " ", regex=False)
        .str.strip()
    )
    return df

def pick_col(df, candidates):
    for c in candidates:
        if c in df.columns:
            return c
    raise KeyError(f"Colonne introuvable. Cherché: {candidates}. Dispo: {list(df.columns)}")

def parse_time(series):
    s = series.astype(str).str.strip()
    s = s.str.replace(r"^(\d):", r"0\1:", regex=True)
    return pd.to_datetime(s, format="%H:%M", errors="coerce")

def safe_sheet_name(name):
    name = str(name) if name else "Inconnu"
    name = re.sub(r'[\\/*?:\[\]/]', '-', name).strip()
    return (name or "Inconnu")[:31]

def fmt_passage(x, i):
    row = x.iloc[i]
    tag = "*" if row.get("Auto", False) else ""
    return f"{row['Heure']} – {row['Métier']} (Table {row['Table']}){tag}"

# ==========================================
# CŒUR DE L'AFFECTATION
# ==========================================

def run_affectation(df_voeux, df_tables, df_groupes, absents: list, melanger_eleves: bool):
    """
    Similaire au code Gradio final : 
    - Mélange aléatoirement les élèves si demandé.
    - Retire les métiers des absents.
    - Retourne df_affectations et df_publipostage.
    """
    df_voeux   = clean_cols(df_voeux.copy())
    df_tables  = clean_cols(df_tables.copy())
    df_groupes = clean_cols(df_groupes.copy())

    # -------- CORRECTION BIAIS: Mélange aléatoire --------
    if melanger_eleves:
        df_voeux = df_voeux.sample(frac=1, random_state=None).reset_index(drop=True)
    # -----------------------------------------------------

    # Colonnes dynamiques
    col_g_debut = pick_col(df_groupes, ["Horaire début","Horaire debut","Heure début","Heure debut"])
    col_g_fin   = pick_col(df_groupes, ["Horaire fin","Horaire Fin","Heure fin","Heure Fin"])
    col_g_grp   = pick_col(df_groupes, ["Groupe","Classe"])
    col_t_metier  = pick_col(df_tables, ["Metier","Métier"])
    col_t_cap     = pick_col(df_tables, ["Capacite par creneau","Capacité par creneau","Capacite","Capacité"])
    col_t_debut   = pick_col(df_tables, ["Heure debut","Heure début","Début","Debut"])
    col_t_fin     = pick_col(df_tables, ["Heure fin","Fin"])
    col_t_interv  = pick_col(df_tables, ["Nom Intervenant","Intervenant","Nom intervenant"])

    # Normalisation nom de colonne Table
    if "Table" not in df_tables.columns and "table" in df_tables.columns:
        df_tables.rename(columns={"table": "Table"}, inplace=True)
    if "Table" not in df_tables.columns:
        df_tables["Table"] = "N/A" # Sécurité si pas de table

    # Retrait des absents
    if absents:
        df_tables = df_tables[~df_tables[col_t_metier].isin(absents)].copy()

    # Parsing heures
    df_groupes["_debut"] = parse_time(df_groupes[col_g_debut])
    df_groupes["_fin"]   = parse_time(df_groupes[col_g_fin])
    df_tables["_debut"]  = parse_time(df_tables[col_t_debut])
    df_tables["_fin"]    = parse_time(df_tables[col_t_fin])

    # Créneaux par groupe
    groupes_horaires = {}
    for _, row in df_groupes.iterrows():
        groupe = row[col_g_grp]
        creneaux, cur = [], row["_debut"]
        while cur + timedelta(minutes=15) <= row["_fin"]:
            creneaux.append(cur.strftime("%H:%M"))
            cur += timedelta(minutes=15)
        groupes_horaires[groupe] = creneaux

    # Agenda capacité par (métier, heure)
    agenda = {}
    for _, row in df_tables.iterrows():
        metier, cap = row[col_t_metier], int(row[col_t_cap])
        cur = row["_debut"]
        while cur + timedelta(minutes=15) <= row["_fin"]:
            agenda[(metier, cur.strftime("%H:%M"))] = cap
            cur += timedelta(minutes=15)

    metiers_disponibles = list(df_tables[col_t_metier].dropna().unique())

    def metiers_par_dispo(agenda, metiers, heure, exclus):
        candidats = []
        for m in metiers:
            if m in exclus:
                continue
            cap = agenda.get((m, heure), 0)
            if cap > 0:
                candidats.append((m, cap))
        candidats.sort(key=lambda x: -x[1])
        return [m for m, _ in candidats]

    # Affectation
    affectations = []
    for _, eleve in df_voeux.iterrows():
        nom, prenom, classe = eleve.get("Nom", "Inconnu"), eleve.get("Prénom", "Inconnu"), eleve.get("Classe", "Inconnu")
        creneaux = groupes_horaires.get(classe, [])
        voeux = [eleve.get(f"Vœu {i}", None) for i in range(1, 7)]
        voeux_valides = [v for v in voeux if v and not pd.isna(v)]
        dejas, affecte_creneaux = set(), set()

        for vœu in voeux_valides:
            for heure in creneaux:
                key = (vœu, heure)
                if (key in agenda and agenda[key] > 0
                        and vœu not in dejas
                        and heure not in affecte_creneaux):
                    affectations.append({
                        "Nom": nom, "Prénom": prenom, "Classe": classe,
                        "Heure": heure, "Métier": vœu,
                        "Vœu n°": voeux.index(vœu) + 1,
                        "Auto": False
                    })
                    agenda[key] -= 1
                    dejas.add(vœu)
                    affecte_creneaux.add(heure)
                    break
            if len(affecte_creneaux) >= 4:
                break

        if len(affecte_creneaux) < 4:
            creneaux_libres = [h for h in creneaux if h not in affecte_creneaux]
            for heure in creneaux_libres:
                if len(affecte_creneaux) >= 4:
                    break
                candidats = metiers_par_dispo(agenda, metiers_disponibles, heure, dejas)
                if candidats:
                    metier_auto = candidats[0]
                    affectations.append({
                        "Nom": nom, "Prénom": prenom, "Classe": classe,
                        "Heure": heure, "Métier": metier_auto,
                        "Vœu n°": 0,
                        "Auto": True
                    })
                    agenda[(metier_auto, heure)] -= 1
                    dejas.add(metier_auto)
                    affecte_creneaux.add(heure)

    df_aff = pd.DataFrame(affectations)
    if df_aff.empty:
        return df_aff, pd.DataFrame(), df_tables, col_t_metier, col_t_interv

    metier_to_interv = dict(zip(df_tables[col_t_metier], df_tables[col_t_interv]))
    metier_to_table  = dict(zip(df_tables[col_t_metier], df_tables["Table"]))
    df_aff["Intervenant"] = df_aff["Métier"].map(metier_to_interv).fillna("Inconnu")
    df_aff["Table"]       = df_aff["Métier"].map(metier_to_table).fillna("N/A")

    df_sorted = df_aff.sort_values(["Classe","Nom","Prénom","Heure"])
    df_pub = (
        df_sorted.groupby(["Nom","Prénom","Classe"])
        .apply(lambda x: pd.Series({
            "Passage 1": fmt_passage(x, 0) if len(x) > 0 else "",
            "Passage 2": fmt_passage(x, 1) if len(x) > 1 else "",
            "Passage 3": fmt_passage(x, 2) if len(x) > 2 else "",
            "Passage 4": fmt_passage(x, 3) if len(x) > 3 else "",
        }))
        .reset_index()
    )

    return df_aff, df_pub, df_tables, col_t_metier, col_t_interv

# ==========================================
# GENERATION DES FICHIERS (IN MEMORY)
# ==========================================

def add_cut_line(doc):
    p = doc.add_paragraph("✂  ─────────────────────────────────────────  ✂")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(9)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    top = OxmlElement('w:top')
    top.set(qn('w:val'), 'single'); top.set(qn('w:sz'), '6')
    top.set(qn('w:space'), '6'); top.set(qn('w:color'), '999999')
    pBdr.append(top); pPr.append(pBdr)

def add_student_block(doc, row, plan_img_bytes=None):
    titre = doc.add_paragraph(f"{row['Prénom']} {row['Nom']}  –  {row['Classe']}")
    titre.runs[0].bold = True
    titre.alignment = WD_ALIGN_PARAGRAPH.CENTER

    t = doc.add_table(rows=1, cols=2)
    cl, cr = t.rows[0].cells[0], t.rows[0].cells[1]

    cl.paragraphs[0].add_run("🕒 Planning").bold = True
    for i in range(1, 5):
        passage = row.get(f"Passage {i}", "")
        if passage:
            para = cl.add_paragraph(passage, style="List Bullet")
            for r in para.runs:
                r.font.size = Pt(10)

    cr.paragraphs[0].add_run("🗺 Plan des tables").bold = True
    if plan_img_bytes:
        cr.add_paragraph()
        run = cr.paragraphs[1].add_run()
        run.add_picture(plan_img_bytes, width=Cm(12.03), height=Cm(7.23))
    else:
        cr.add_paragraph("(plan non fourni)")
    doc.add_paragraph("")

def gen_word_eleves(df_pub, plan_img_bytes=None):
    doc = Document()
    for section in doc.sections:
        section.top_margin    = Cm(0.8)
        section.bottom_margin = Cm(0.8)
        section.left_margin   = Cm(1.0)
        section.right_margin  = Cm(1.0)

    doc.add_heading("Planning individuel – Forum Orientation", 0)
    rows = df_pub.sort_values(["Classe", "Nom", "Prénom"]).to_dict(orient="records")
    for idx in range(0, len(rows), 2):
        if idx > 0:
            doc.add_page_break()
        add_student_block(doc, rows[idx], plan_img_bytes)
        add_cut_line(doc)
        if idx + 1 < len(rows):
            add_student_block(doc, rows[idx + 1], plan_img_bytes)
            
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

def gen_word_intervenants(df_aff):
    doc = Document()
    doc.add_heading("Planning des Intervenants – Forum Orientation", 0)
    first = True
    for interv, grp in df_aff.sort_values(["Intervenant","Heure"]).groupby("Intervenant"):
        if not first:
            doc.add_page_break()
        first = False
        doc.add_heading(f"Planning de {interv}", level=1)
        cur_heure = None
        for _, row in grp.iterrows():
            if row["Heure"] != cur_heure:
                doc.add_paragraph(f"⏰  {row['Heure']}")
                cur_heure = row["Heure"]
            doc.add_paragraph(f"  {row['Nom']} {row['Prénom']} ({row['Classe']}) – {row['Métier']}")
            
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

def gen_excel_publipostage(df_pub):
    bio = io.BytesIO()
    with pd.ExcelWriter(bio, engine='openpyxl') as writer:
        df_pub.to_excel(writer, index=False)
    return bio.getvalue()

def gen_excel_intervenants(df_aff):
    bio = io.BytesIO()
    with pd.ExcelWriter(bio, engine='openpyxl') as writer:
        for interv, grp in df_aff.sort_values(["Intervenant","Heure","Nom"]).groupby("Intervenant"):
            grp[["Intervenant","Table","Métier","Heure","Nom","Prénom","Classe"]].to_excel(
                writer, sheet_name=safe_sheet_name(interv), index=False
            )
    return bio.getvalue()

# ==========================================
# INTERFACE STREAMLIT
# ==========================================

st.title("🎓 Forum Orientation – Gestion des affectations")
st.markdown("Bienvenue dans l'application de répartition automatique des élèves aux tables des professionnels.")

st.header("1️⃣ Charger les fichiers")
col1, col2, col3 = st.columns(3)
file_voeux = col1.file_uploader("📋 Vœux élèves (Excel)", type=["xlsx"])
file_tables = col2.file_uploader("🏷 Tables & pôles (Excel)", type=["xlsx"])
file_groupes = col3.file_uploader("🕐 Groupes horaires (Excel)", type=["xlsx"])

file_plan = st.file_uploader("🗺 Plan de salle (image PNG/JPG – optionnel)", type=["png", "jpg", "jpeg"])

if file_voeux and file_tables and file_groupes:
    # On charge justes les tables pour extraire les métiers absents potentiels
    try:
        df_t = pd.read_excel(file_tables)
        df_t_clean = clean_cols(df_t)
        col_t_metier = pick_col(df_t_clean, ["Metier","Métier"])
        metiers_list = sorted(df_t_clean[col_t_metier].dropna().unique().tolist())
        
        st.success(f"✅ Fichiers chargés – {len(metiers_list)} métiers trouvés.")
        
        st.header("2️⃣ Paramètres de l'affectation")
        absents = st.multiselect("⚠️ Métiers absents le jour J", options=metiers_list, help="Sélectionnez les métiers dont l'intervenant est absent. Ils ne seront pas affectés.")
        melanger_eleves = st.checkbox("🎲 Mélanger aléatoirement l'ordre des élèves pour plus d'équité", value=True, help="Corrige le biais où les élèves en fin de fichier obtiennent moins de leurs vœux initiaux.")
        
        if st.button("🚀 Lancer la génération des plannings (ZIP)"):
            with st.spinner("Affectation en cours et génération des documents..."):
                try:
                    df_voeux = pd.read_excel(file_voeux)
                    file_tables.seek(0)
                    df_tables = pd.read_excel(file_tables)
                    df_groupes = pd.read_excel(file_groupes)
                    
                    df_aff, df_pub, _, _, _ = run_affectation(df_voeux, df_tables, df_groupes, absents, melanger_eleves)
                    
                    if df_aff.empty:
                        st.error("Aucune affectation générée. Vérifiez vos fichiers.")
                    else:
                        # Generation des fichiers en mémoire
                        plan_bytes = io.BytesIO(file_plan.read()) if file_plan else None
                        
                        xl_pub = gen_excel_publipostage(df_pub)
                        xl_int = gen_excel_intervenants(df_aff)
                        wd_eleves = gen_word_eleves(df_pub, plan_bytes)
                        wd_int = gen_word_intervenants(df_aff)
                        
                        # Creation du ZIP en memoire
                        zip_buffer = io.BytesIO()
                        with zipfile.ZipFile(zip_buffer, "w") as z:
                            z.writestr("planning_publipostage.xlsx", xl_pub)
                            z.writestr("planning_par_intervenant.xlsx", xl_int)
                            z.writestr("planning_eleves.docx", wd_eleves)
                            z.writestr("planning_intervenants.docx", wd_int)
                        
                        # Stats
                        nb_eleves = len(df_pub)
                        nb_aff    = len(df_aff)
                        nb_auto   = int(df_aff["Auto"].sum()) if "Auto" in df_aff.columns else 0
                        nb_voeux  = nb_aff - nb_auto
                        
                        st.success("✅ Plannings générés avec succès !")
                        st.info(f"📊 **Statistiques :**\n- **{nb_eleves}** élèves traités\n- **{nb_voeux}** passages sur vœux\n- **{nb_auto}** passages suggérés automatiquement")
                        
                        st.download_button(
                            label="📦 Télécharger le ZIP (tous les documents)",
                            data=zip_buffer.getvalue(),
                            file_name="forum_orientation_exports.zip",
                            mime="application/zip",
                            type="primary"
                        )
                except Exception as e:
                    import traceback
                    st.error(f"Une erreur est survenue durant la génération :\n{e}")
                    with st.expander("Voir les détails techniques (Traceback)"):
                        st.code(traceback.format_exc())
                        
    except Exception as e:
        st.error(f"Erreur lors de la lecture des en-têtes du fichier Excel Tables : {e}")
else:
    st.info("Veuillez charger les 3 fichiers Excel (Vœux, Tables, Groupes) pour démarrer.")